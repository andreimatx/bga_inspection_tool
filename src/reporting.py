"""IPC-aligned CSV, Excel and Word reporting for BGA void inspection.

Every inspection produces, per board:
- ``<stem>_ball_metrics.csv``      raw per-ball table (machine friendly)
- ``<stem>_summary.csv``           one-line board summary (machine friendly)
- ``<stem>_inspection_report.xlsx``presentation-grade workbook, 4 sheets:
      1. Executive Summary  - verdict banner, key metrics, distribution
      2. IPC Evaluation     - each criterion, its source, measured, result
      3. Ball Details       - full table with conditional formatting
      4. Algorithm Analysis - honest description of how the numbers were
                              produced, uncertainty and manual-check list
- ``<stem>_inspection_report.docx``printable Word report with charts and
                              the annotated X-ray preview embedded.

Acceptance criteria follow IPC-A-610 (25% cumulative projected void area per
ball = defect) with IPC-7095 process guidance (10-25% band = process
indicator). See ``src/standards.py`` for the exact numbers and caveats.
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd

from src.standards import (
    IPCVoidCriteria,
    evaluate_board,
    ipc_assessment_label,
    void_diameter_ratio_percent,
)

# ---------------------------------------------------------------------------
# Shared palette (hex, no leading #) so Excel and Word match visually.
# ---------------------------------------------------------------------------
COLOR_HEADER_BG = "1F3864"      # dark navy
COLOR_HEADER_FG = "FFFFFF"
COLOR_ACCENT_BG = "D6E4F0"      # light steel blue
COLOR_PASS = "2E7D32"           # green
COLOR_PASS_LIGHT = "E3F0E3"
COLOR_REVIEW = "B26A00"         # amber (dark enough for white text)
COLOR_REVIEW_LIGHT = "FDF3D9"
COLOR_FAIL = "C62828"           # red
COLOR_FAIL_LIGHT = "FBE4E4"
COLOR_NOTE = "666666"

TOOL_NAME = "BGA X-ray Void Inspection Tool"


# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------
def build_summary(
    image_name: str,
    rows: list[dict[str, object]],
    warning_threshold: float,
) -> dict[str, object]:
    """Build one image-level statistical summary (CSV-friendly flat dict)."""
    criteria = IPCVoidCriteria()

    if not rows:
        return {
            "image_name": image_name,
            "total_balls": 0,
            "average_void_ratio_percent": 0.0,
            "maximum_void_ratio_percent": 0.0,
            "median_void_ratio_percent": 0.0,
            "max_single_void_ratio_percent": 0.0,
            "balls_over_warning_threshold": 0,
            "warning_threshold_percent": warning_threshold,
            "ipc_a610_defect_balls": 0,
            "ipc_7095_process_indicator_balls": 0,
            "ipc_verdict": "NO BALLS DETECTED",
            "standards_applied": criteria.standard_names,
        }

    ratios = sorted(float(row["void_ratio_percent"]) for row in rows)
    board = evaluate_board(rows, criteria)
    median = ratios[len(ratios) // 2]

    return {
        "image_name": image_name,
        "total_balls": len(rows),
        "average_void_ratio_percent": round(sum(ratios) / len(ratios), 4),
        "maximum_void_ratio_percent": round(max(ratios), 4),
        "median_void_ratio_percent": round(median, 4),
        "max_single_void_ratio_percent": round(
            float(board["max_single_void_percent"]), 4
        ),
        "balls_over_warning_threshold": sum(
            ratio >= warning_threshold for ratio in ratios
        ),
        "warning_threshold_percent": warning_threshold,
        "ipc_a610_defect_balls": board["defect_balls"],
        "ipc_7095_process_indicator_balls": board["process_indicator_balls"],
        "ipc_verdict": board["ipc_verdict"],
        "standards_applied": criteria.standard_names,
    }


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
def save_reports(
    rows: list[dict[str, object]],
    summary: dict[str, object],
    output_dir: Path,
    image_stem: str,
    context: dict[str, object] | None = None,
) -> dict[str, Path]:
    """Save CSV, Excel and Word reports. ``context`` carries extra run info.

    Recognized context keys (all optional):
        image_path, image_width, image_height, expected_slots,
        occupied_slots, warning_threshold, fail_threshold,
        largest_void_review_threshold, preview_image, contours_image,
        processing_seconds
    """
    context = dict(context or {})
    criteria = IPCVoidCriteria()
    enriched_rows = _enrich_rows(rows, criteria)
    board = evaluate_board(rows, criteria) if rows else None

    details_df = pd.DataFrame(enriched_rows)
    summary_df = pd.DataFrame([summary])

    details_csv = output_dir / f"{image_stem}_ball_metrics.csv"
    summary_csv = output_dir / f"{image_stem}_summary.csv"
    excel_path = output_dir / f"{image_stem}_inspection_report.xlsx"
    word_path = output_dir / f"{image_stem}_inspection_report.docx"

    details_df.to_csv(details_csv, index=False)
    summary_df.to_csv(summary_csv, index=False)

    _write_excel_report(
        excel_path, enriched_rows, summary, board, criteria, image_stem, context
    )
    _write_word_report(
        word_path, enriched_rows, summary, board, criteria, image_stem, context
    )

    return {
        "details_csv": details_csv,
        "summary_csv": summary_csv,
        "excel": excel_path,
        "word": word_path,
    }


def _enrich_rows(
    rows: list[dict[str, object]],
    criteria: IPCVoidCriteria,
) -> list[dict[str, object]]:
    """Append derived, standards-facing columns to each per-ball row."""
    enriched = []
    for row in rows:
        new_row = dict(row)
        ball_area = float(row.get("ball_area_px", 0.0) or 0.0)
        largest = float(row.get("largest_void_area_px", 0.0) or 0.0)
        new_row["largest_void_diameter_percent"] = round(
            void_diameter_ratio_percent(largest, ball_area), 2
        )
        new_row["ipc_assessment"] = ipc_assessment_label(
            str(row.get("quality", ""))
        )
        enriched.append(new_row)
    return enriched


def _quality_counts(rows: list[dict[str, object]]) -> dict[str, int]:
    counts = {"PASS": 0, "REVIEW": 0, "FAIL": 0}
    for row in rows:
        quality = str(row.get("quality", ""))
        if quality in counts:
            counts[quality] += 1
    return counts


def _verdict_color(verdict: str) -> tuple[str, str]:
    """Return (solid color, light fill) for a board verdict string."""
    if "REJECT" in verdict or "FAIL" in verdict:
        return COLOR_FAIL, COLOR_FAIL_LIGHT
    if "INDICATOR" in verdict or "REVIEW" in verdict:
        return COLOR_REVIEW, COLOR_REVIEW_LIGHT
    return COLOR_PASS, COLOR_PASS_LIGHT


# ---------------------------------------------------------------------------
# Criterion evaluation rows shared by Excel and Word
# ---------------------------------------------------------------------------
def _criteria_rows(
    summary: dict[str, object],
    board: dict[str, object] | None,
    criteria: IPCVoidCriteria,
    context: dict[str, object],
) -> list[tuple[str, str, str, str, str]]:
    """(criterion, source, limit, measured, result) rows for the reports."""
    if board is None:
        return []

    max_total = float(board["max_total_void_percent"])
    max_single_dia = float(board["max_single_void_diameter_percent"])
    defect_balls = int(board["defect_balls"])
    indicator_balls = int(board["process_indicator_balls"])
    dominant = int(board["dominant_void_balls"])

    expected = context.get("expected_slots")
    detected = int(summary.get("total_balls", 0))
    detection_measured = (
        f"{detected} / {expected}" if expected else f"{detected}"
    )
    detection_ok = (expected is None) or (detected == int(expected))

    return [
        (
            "Cumulative void area per ball",
            "IPC-A-610 (Class 1, 2, 3)",
            f"<= {criteria.defect_area_percent:.0f}% of ball area",
            f"worst ball: {max_total:.2f}%",
            "PASS" if defect_balls == 0 else f"FAIL ({defect_balls} ball(s))",
        ),
        (
            "Process indicator band (total void 10-25%)",
            "IPC-7095 process guidance",
            "monitor, not a defect",
            f"{indicator_balls} ball(s) in band",
            "INFO" if indicator_balls else "PASS",
        ),
        (
            "Largest single void per ball",
            "IPC-7095 (dominant-void guidance)",
            (
                f"review above {criteria.largest_void_review_area_percent:.1f}% "
                "area (~35% diameter)"
            ),
            f"worst single void: {max_single_dia:.1f}% of ball diameter",
            "INFO" if dominant else "PASS",
        ),
        (
            "Void location at pad interface",
            "IPC-7095 (30% dia at land)",
            "requires oblique/3D view",
            "not measurable in one top-down 2D X-ray",
            "NOT EVALUATED",
        ),
        (
            "Complete ball grid detected",
            "detection integrity gate",
            "all expected pads found",
            detection_measured,
            "PASS" if detection_ok else "REVIEW",
        ),
    ]


def _honest_analysis_sections(
    rows: list[dict[str, object]],
    summary: dict[str, object],
    context: dict[str, object],
) -> list[tuple[str, list[str]]]:
    """The 'no-excuses' algorithm analysis content, shared by Excel/Word."""
    counts = _quality_counts(rows)
    estimated = [r for r in rows if bool(r.get("is_estimated_pad"))]
    confidences = [
        float(r.get("confidence", 0.0) or 0.0)
        for r in rows
        if not bool(r.get("is_estimated_pad"))
    ]
    mean_conf = sum(confidences) / len(confidences) if confidences else 0.0

    top = sorted(
        rows,
        key=lambda r: float(r.get("void_ratio_percent", 0.0)),
        reverse=True,
    )[:10]
    top_list = [
        (
            f"Ball #{r['ball_id']:>3} at ({r['center_x']}, {r['center_y']}): "
            f"total {float(r['void_ratio_percent']):.2f}%, "
            f"largest single {float(r['largest_void_ratio_percent']):.2f}% "
            f"-> {r['ipc_assessment']}"
        )
        for r in top
    ]

    est_list = [
        (
            f"Ball #{r['ball_id']:>3} at ({r['center_x']}, {r['center_y']}) "
            "- position estimated from the grid, no direct dark-pad evidence; "
            "its void metrics carry extra uncertainty"
        )
        for r in estimated
    ] or ["None - every pad was detected from direct image evidence."]

    pipeline = [
        "1. CLAHE contrast enhancement + median denoising of the raw X-ray.",
        "2. Solder-ball detection: Hough circles + dark-blob evidence, "
        "regularized onto the 16x16 BGA grid (affine + homography fit for "
        "oblique views, per-ball local snap with non-degradation gates).",
        "3. Per-ball void segmentation: median-background subtraction "
        "(kernel > ball radius), bond-wire suppression by grayscale "
        "closing, residual threshold, then shape filtering.",
        "4. Detected void blobs are rendered as fitted circles when they "
        "fill >= 45% of the fitted disc (physical voids are spherical gas "
        "bubbles); ragged non-circular structures keep their raw contour.",
        "5. Component-level sanity filters remove rim arcs, trace stripes, "
        "bare-board bleed-through and empty-pad artifacts before metrics.",
    ]

    uncertainty = [
        "Projected-area measurement: one top-down 2D X-ray cannot separate "
        "package-interface, mid-ball and pad-interface voids; the 25% "
        "IPC-A-610 criterion is applied on the projection, exactly as the "
        "standard defines it for X-ray inspection.",
        "Analysis is restricted to the inner 82% of each ball radius; a "
        "void hiding exactly on the ball rim can be underestimated.",
        "Voids smaller than ~10 px^2 (about 0.15% of a ball) are below the "
        "detection floor and are intentionally ignored as noise.",
        "Fitted-circle rendering slightly increases void area versus the "
        "raw threshold blob (it includes the faded bubble edge); this is "
        "closer to the physical bubble size but adds ~0.3-0.6 percentage "
        "points to board averages compared to raw-contour reporting.",
        "JPEG compression artifacts and CLAHE amplification set a noise "
        "floor around 1-2% void ratio; differences below that level are "
        "not meaningful.",
        "On oblique views, perspective foreshortening distorts both ball "
        "and void areas; the ratio is partially self-correcting but not "
        "perfectly.",
    ]

    limitations = [
        "Components mounted on the opposite board side project dark "
        "shadows between and over balls; the detector anchors circles on "
        "grid geometry to resist this, but heavily shadowed balls have "
        "reduced void sensitivity.",
        "Grid slots occupied by vias (no ball populated) are still "
        "reported as balls when the slot must be filled; they show ~0% "
        "void and near-board brightness.",
        "Bond-wire suppression (morphological closing) can mask very thin "
        "elongated voids lying exactly under a wire.",
        "The tool measures voiding only; it does not assess head-in-pillow, "
        "non-wet opens, bridging or solder-ball shape defects.",
    ]

    thresholds_used = [
        f"Defect (FAIL): total void >= "
        f"{context.get('fail_threshold', 25.0)}% of ball area (IPC-A-610).",
        f"Process indicator (REVIEW): total void >= "
        f"{context.get('warning_threshold', 10.0)}% (IPC-7095 guidance).",
        "Dominant single void (REVIEW): largest void >= "
        f"{context.get('largest_void_review_threshold', 12.25)}% of ball "
        "area (~35% of ball diameter).",
    ]

    detection_facts = [
        f"Balls detected: {summary.get('total_balls', 0)}"
        + (
            f" of {context['expected_slots']} expected grid slots"
            if context.get("expected_slots")
            else ""
        ),
        f"Pads placed from grid estimation (no direct evidence): "
        f"{len(estimated)}",
        f"Mean detection confidence (evidence-backed pads): {mean_conf:.2f}",
        f"Ball quality split: {counts['PASS']} PASS / {counts['REVIEW']} "
        f"REVIEW / {counts['FAIL']} FAIL",
    ]

    return [
        ("Detection integrity", detection_facts),
        ("Measurement pipeline (what actually produced these numbers)", pipeline),
        ("Classification thresholds applied", thresholds_used),
        ("Measurement uncertainty - read before trusting decimals", uncertainty),
        ("Known limitations of this analysis", limitations),
        ("Balls to verify manually first (worst 10 by total void)", top_list),
        ("Grid-estimated pads (verify visually)", est_list),
    ]


# ---------------------------------------------------------------------------
# Excel report
# ---------------------------------------------------------------------------
def _write_excel_report(
    path: Path,
    rows: list[dict[str, object]],
    summary: dict[str, object],
    board: dict[str, object] | None,
    criteria: IPCVoidCriteria,
    image_stem: str,
    context: dict[str, object],
) -> None:
    import xlsxwriter

    workbook = xlsxwriter.Workbook(str(path), {"nan_inf_to_errors": True})

    fmt_title = workbook.add_format(
        {
            "bold": True,
            "font_size": 18,
            "font_color": COLOR_HEADER_FG,
            "bg_color": COLOR_HEADER_BG,
            "align": "center",
            "valign": "vcenter",
        }
    )
    fmt_subtitle = workbook.add_format(
        {
            "italic": True,
            "font_size": 10,
            "font_color": COLOR_HEADER_FG,
            "bg_color": COLOR_HEADER_BG,
            "align": "center",
            "valign": "vcenter",
        }
    )
    fmt_section = workbook.add_format(
        {
            "bold": True,
            "font_size": 12,
            "font_color": COLOR_HEADER_BG,
            "bottom": 2,
            "border_color": COLOR_HEADER_BG,
        }
    )
    fmt_label = workbook.add_format(
        {"bold": True, "bg_color": COLOR_ACCENT_BG, "border": 1}
    )
    fmt_value = workbook.add_format({"border": 1})
    fmt_value_num = workbook.add_format({"border": 1, "num_format": "0.00"})
    fmt_header = workbook.add_format(
        {
            "bold": True,
            "font_color": COLOR_HEADER_FG,
            "bg_color": COLOR_HEADER_BG,
            "border": 1,
            "align": "center",
            "valign": "vcenter",
            "text_wrap": True,
        }
    )
    fmt_cell = workbook.add_format({"border": 1})
    fmt_cell_num = workbook.add_format({"border": 1, "num_format": "0.00"})
    fmt_note = workbook.add_format(
        {"italic": True, "font_color": COLOR_NOTE, "text_wrap": True}
    )
    fmt_wrap = workbook.add_format({"border": 1, "text_wrap": True, "valign": "top"})

    def verdict_formats(verdict: str):
        solid, light = _verdict_color(verdict)
        big = workbook.add_format(
            {
                "bold": True,
                "font_size": 16,
                "font_color": "FFFFFF",
                "bg_color": solid,
                "align": "center",
                "valign": "vcenter",
                "border": 1,
            }
        )
        detail = workbook.add_format(
            {
                "italic": True,
                "bg_color": light,
                "align": "center",
                "valign": "vcenter",
                "text_wrap": True,
                "border": 1,
            }
        )
        return big, detail

    result_formats = {
        "PASS": workbook.add_format(
            {"bold": True, "font_color": "FFFFFF", "bg_color": COLOR_PASS,
             "align": "center", "border": 1}
        ),
        "FAIL": workbook.add_format(
            {"bold": True, "font_color": "FFFFFF", "bg_color": COLOR_FAIL,
             "align": "center", "border": 1}
        ),
        "INFO": workbook.add_format(
            {"bold": True, "font_color": "FFFFFF", "bg_color": COLOR_REVIEW,
             "align": "center", "border": 1}
        ),
        "REVIEW": workbook.add_format(
            {"bold": True, "font_color": "FFFFFF", "bg_color": COLOR_REVIEW,
             "align": "center", "border": 1}
        ),
        "NOT EVALUATED": workbook.add_format(
            {"bold": True, "font_color": COLOR_NOTE, "bg_color": "EEEEEE",
             "align": "center", "border": 1}
        ),
    }

    generated = datetime.now().strftime("%Y-%m-%d %H:%M")
    verdict = str(summary.get("ipc_verdict", "N/A"))

    # ----------------------------- Sheet 1: Executive Summary -------------
    ws = workbook.add_worksheet("Executive Summary")
    ws.hide_gridlines(2)
    ws.set_column("A:A", 2)
    ws.set_column("B:C", 26)
    ws.set_column("D:E", 22)
    ws.set_column("F:G", 16)

    ws.merge_range("B2:G3", "BGA VOID INSPECTION REPORT", fmt_title)
    ws.merge_range(
        "B4:G4",
        f"Quantitative & qualitative void analysis per {criteria.standard_names}",
        fmt_subtitle,
    )

    row_i = 6
    ws.write(row_i, 1, "Inspection details", fmt_section)
    row_i += 1
    info_pairs = [
        ("Board / image", str(summary.get("image_name", image_stem))),
        ("Report generated", generated),
        ("Inspection tool", TOOL_NAME),
        (
            "Image size",
            f"{context.get('image_width', '?')} x "
            f"{context.get('image_height', '?')} px",
        ),
        (
            "BGA grid",
            f"{context.get('expected_slots', '?')} expected pads (16 x 16)",
        ),
        ("Balls analyzed", str(summary.get("total_balls", 0))),
    ]
    if context.get("processing_seconds"):
        info_pairs.append(
            ("Processing time", f"{float(context['processing_seconds']):.1f} s")
        )
    for label, value in info_pairs:
        ws.write(row_i, 1, label, fmt_label)
        ws.write(row_i, 2, value, fmt_value)
        row_i += 1

    # Verdict banner.
    fmt_big, fmt_detail = verdict_formats(verdict)
    ws.merge_range(6, 3, 8, 6, f"IPC VERDICT:  {verdict}", fmt_big)
    detail = str(board["ipc_verdict_detail"]) if board else ""
    ws.merge_range(9, 3, 11, 6, detail, fmt_detail)

    # Key metrics.
    metrics_row = row_i + 1
    ws.write(metrics_row, 1, "Key void metrics", fmt_section)
    metrics_row += 1
    counts = _quality_counts(rows)
    key_metrics = [
        ("Average void ratio", float(summary.get("average_void_ratio_percent", 0.0)), "%"),
        ("Median void ratio", float(summary.get("median_void_ratio_percent", 0.0)), "%"),
        ("Worst ball (total void)", float(summary.get("maximum_void_ratio_percent", 0.0)), "%"),
        ("Worst single void", float(summary.get("max_single_void_ratio_percent", 0.0)), "%"),
        ("Balls ACCEPTABLE", counts["PASS"], "balls"),
        ("Balls PROCESS INDICATOR", counts["REVIEW"], "balls"),
        ("Balls DEFECT (IPC-A-610)", counts["FAIL"], "balls"),
    ]
    for label, value, unit in key_metrics:
        ws.write(metrics_row, 1, label, fmt_label)
        if isinstance(value, float):
            ws.write_number(metrics_row, 2, value, fmt_value_num)
        else:
            ws.write_number(metrics_row, 2, value, fmt_value)
        ws.write(metrics_row, 3, unit, fmt_value)
        metrics_row += 1

    # Distribution table + charts.
    dist_row = metrics_row + 2
    ws.write(dist_row, 1, "Void-ratio distribution", fmt_section)
    dist_row += 1
    ws.write(dist_row, 1, "Total void per ball", fmt_header)
    ws.write(dist_row, 2, "Ball count", fmt_header)
    distribution = board["distribution"] if board else []
    for offset, (label, count) in enumerate(distribution, start=1):
        ws.write(dist_row + offset, 1, label, fmt_cell)
        ws.write_number(dist_row + offset, 2, count, fmt_cell)

    if distribution:
        chart = workbook.add_chart({"type": "column"})
        n = len(distribution)
        chart.add_series(
            {
                "name": "Balls per void band",
                "categories": ["Executive Summary", dist_row + 1, 1, dist_row + n, 1],
                "values": ["Executive Summary", dist_row + 1, 2, dist_row + n, 2],
                "fill": {"color": "#4472C4"},
                "data_labels": {"value": True},
            }
        )
        chart.set_title({"name": "Void-ratio distribution (per ball)"})
        chart.set_x_axis({"name": "Total void per ball"})
        chart.set_y_axis({"name": "Number of balls", "major_gridlines": {"visible": False}})
        chart.set_legend({"none": True})
        chart.set_size({"width": 520, "height": 300})
        ws.insert_chart(dist_row, 4, chart)

        pie_start = dist_row + n + 2
        ws.write(pie_start, 1, "Assessment", fmt_header)
        ws.write(pie_start, 2, "Balls", fmt_header)
        pie_data = [
            ("ACCEPTABLE", counts["PASS"]),
            ("PROCESS INDICATOR", counts["REVIEW"]),
            ("DEFECT", counts["FAIL"]),
        ]
        for offset, (label, count) in enumerate(pie_data, start=1):
            ws.write(pie_start + offset, 1, label, fmt_cell)
            ws.write_number(pie_start + offset, 2, count, fmt_cell)
        pie = workbook.add_chart({"type": "doughnut"})
        pie.add_series(
            {
                "name": "Ball assessment",
                "categories": ["Executive Summary", pie_start + 1, 1, pie_start + 3, 1],
                "values": ["Executive Summary", pie_start + 1, 2, pie_start + 3, 2],
                "points": [
                    {"fill": {"color": "#" + COLOR_PASS}},
                    {"fill": {"color": "#" + COLOR_REVIEW}},
                    {"fill": {"color": "#" + COLOR_FAIL}},
                ],
            }
        )
        pie.set_title({"name": "Ball assessment split"})
        pie.set_size({"width": 360, "height": 280})
        ws.insert_chart(dist_row + 16, 4, pie)

    # ----------------------------- Sheet 2: IPC Evaluation ----------------
    ws2 = workbook.add_worksheet("IPC Evaluation")
    ws2.hide_gridlines(2)
    ws2.set_column("A:A", 2)
    ws2.set_column("B:B", 38)
    ws2.set_column("C:C", 28)
    ws2.set_column("D:D", 34)
    ws2.set_column("E:E", 40)
    ws2.set_column("F:F", 18)

    ws2.merge_range("B2:F2", "IPC CRITERIA EVALUATION", fmt_title)
    headers = ["Criterion", "Source", "Limit", "Measured on this board", "Result"]
    for col, header in enumerate(headers, start=1):
        ws2.write(3, col, header, fmt_header)
    for idx, (crit, source, limit, measured, result) in enumerate(
        _criteria_rows(summary, board, criteria, context), start=4
    ):
        ws2.write(idx, 1, crit, fmt_wrap)
        ws2.write(idx, 2, source, fmt_wrap)
        ws2.write(idx, 3, limit, fmt_wrap)
        ws2.write(idx, 4, measured, fmt_wrap)
        key = result.split(" ")[0] if result.startswith("FAIL") else result
        ws2.write(idx, 5, result, result_formats.get(key, fmt_cell))

    note_row = 11
    ws2.merge_range(
        note_row,
        1,
        note_row + 3,
        5,
        "Standards note: IPC-A-610 defines BGA voiding as a DEFECT for Class "
        "1, 2 and 3 when the cumulative void area of a ball exceeds 25% of "
        "its X-ray image area. IPC-7095 adds process-level guidance: it "
        "classifies void types (macrovoids, planar microvoids, shrinkage, "
        "via-in-pad) and treats the 10-25% band as a process indicator to "
        "monitor. Location-specific limits (e.g. voids at the pad interface) "
        "cannot be judged from a single top-down 2D X-ray and are marked "
        "NOT EVALUATED instead of being guessed.",
        fmt_note,
    )

    # ----------------------------- Sheet 3: Ball Details ------------------
    ws3 = workbook.add_worksheet("Ball Details")
    columns = [
        ("ball_id", "Ball ID", 8),
        ("center_x", "X (px)", 8),
        ("center_y", "Y (px)", 8),
        ("diameter_px", "Diameter (px)", 11),
        ("ball_area_px", "Ball area (px2)", 12),
        ("void_count", "Void count", 9),
        ("total_void_area_px", "Total void area (px2)", 13),
        ("largest_void_area_px", "Largest void (px2)", 13),
        ("void_ratio_percent", "Total void (%)", 11),
        ("largest_void_ratio_percent", "Largest void (%)", 12),
        ("largest_void_diameter_percent", "Largest void dia (% of ball dia)", 15),
        ("quality", "Quality", 9),
        ("ipc_assessment", "IPC assessment", 20),
        ("is_estimated_pad", "Grid-estimated", 12),
        ("confidence", "Confidence", 10),
    ]
    for col, (_, header, width) in enumerate(columns):
        ws3.set_column(col, col, width)
        ws3.write(0, col, header, fmt_header)
    ws3.freeze_panes(1, 0)
    ws3.autofilter(0, 0, max(len(rows), 1), len(columns) - 1)

    fmt_row_fail = workbook.add_format({"border": 1, "bg_color": COLOR_FAIL_LIGHT})
    fmt_row_review = workbook.add_format({"border": 1, "bg_color": COLOR_REVIEW_LIGHT})
    fmt_row_fail_num = workbook.add_format(
        {"border": 1, "bg_color": COLOR_FAIL_LIGHT, "num_format": "0.00"}
    )
    fmt_row_review_num = workbook.add_format(
        {"border": 1, "bg_color": COLOR_REVIEW_LIGHT, "num_format": "0.00"}
    )

    for r_idx, row in enumerate(rows, start=1):
        quality = str(row.get("quality", ""))
        if quality == "FAIL":
            fmt_plain, fmt_num = fmt_row_fail, fmt_row_fail_num
        elif quality == "REVIEW":
            fmt_plain, fmt_num = fmt_row_review, fmt_row_review_num
        else:
            fmt_plain, fmt_num = fmt_cell, fmt_cell_num
        for c_idx, (key, _, _) in enumerate(columns):
            value = row.get(key, "")
            if isinstance(value, bool):
                ws3.write(r_idx, c_idx, "yes" if value else "", fmt_plain)
            elif isinstance(value, (int, float)):
                ws3.write_number(r_idx, c_idx, float(value), fmt_num)
            else:
                ws3.write(r_idx, c_idx, str(value), fmt_plain)

    if rows:
        void_col = 8  # "Total void (%)"
        ws3.conditional_format(
            1,
            void_col,
            len(rows),
            void_col,
            {
                "type": "3_color_scale",
                "min_color": "#E3F0E3",
                "mid_color": "#FDF3D9",
                "max_color": "#F4B5B5",
            },
        )

    # ----------------------------- Sheet 4: Algorithm Analysis ------------
    ws4 = workbook.add_worksheet("Algorithm Analysis")
    ws4.hide_gridlines(2)
    ws4.set_column("A:A", 2)
    ws4.set_column("B:B", 110)
    ws4.merge_range("B2:B3", "", fmt_title)
    ws4.write("B2", "ALGORITHM ANALYSIS - HONEST ASSESSMENT", fmt_title)

    out_row = 4
    for section, bullets in _honest_analysis_sections(rows, summary, context):
        ws4.write(out_row, 1, section, fmt_section)
        out_row += 1
        for bullet in bullets:
            ws4.write(out_row, 1, "  - " + bullet, fmt_note)
            out_row += 1
        out_row += 1

    workbook.close()


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------
def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _histogram_png(rows: list[dict[str, object]]) -> BytesIO | None:
    """Render the void distribution histogram to an in-memory PNG."""
    if not rows:
        return None
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from src.standards import bucket_void_distribution

    ratios = [float(r.get("void_ratio_percent", 0.0)) for r in rows]
    distribution = bucket_void_distribution(ratios)
    labels = [label for label, _ in distribution]
    values = [count for _, count in distribution]
    colors = ["#2E7D32", "#66A96A", "#C9B458", "#E8A33D", "#D66A2C", "#C62828"]

    fig, ax = plt.subplots(figsize=(6.4, 3.2), dpi=150)
    bars = ax.bar(labels, values, color=colors[: len(labels)])
    ax.bar_label(bars, padding=2, fontsize=8)
    ax.set_ylabel("Number of balls")
    ax.set_title("Void-ratio distribution (total void per ball)")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    buffer = BytesIO()
    fig.savefig(buffer, format="png")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _write_word_report(
    path: Path,
    rows: list[dict[str, object]],
    summary: dict[str, object],
    board: dict[str, object] | None,
    criteria: IPCVoidCriteria,
    image_stem: str,
    context: dict[str, object],
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = Document()

    # Page title.
    title = document.add_heading("BGA Void Inspection Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        f"Quantitative & qualitative void analysis per {criteria.standard_names}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    # Info table.
    generated = datetime.now().strftime("%Y-%m-%d %H:%M")
    info_pairs = [
        ("Board / image", str(summary.get("image_name", image_stem))),
        ("Report generated", generated),
        ("Inspection tool", TOOL_NAME),
        ("Standards applied", criteria.standard_names),
        (
            "BGA grid / balls analyzed",
            f"{context.get('expected_slots', '?')} expected pads, "
            f"{summary.get('total_balls', 0)} analyzed",
        ),
    ]
    table = document.add_table(rows=len(info_pairs), cols=2)
    table.style = "Light Grid Accent 1"
    for idx, (label, value) in enumerate(info_pairs):
        cells = table.rows[idx].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value

    # Verdict banner.
    verdict = str(summary.get("ipc_verdict", "N/A"))
    solid, _light = _verdict_color(verdict)
    banner = document.add_table(rows=1, cols=1)
    banner_cell = banner.rows[0].cells[0]
    banner_para = banner_cell.paragraphs[0]
    banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner_para.add_run(f"IPC VERDICT: {verdict}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _shade_cell(banner_cell, solid)
    if board:
        detail = document.add_paragraph(str(board["ipc_verdict_detail"]))
        detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        detail.runs[0].italic = True

    # Key metrics.
    document.add_heading("Key void metrics", level=1)
    counts = _quality_counts(rows)
    metric_pairs = [
        ("Average void ratio", f"{summary.get('average_void_ratio_percent', 0.0)} %"),
        ("Median void ratio", f"{summary.get('median_void_ratio_percent', 0.0)} %"),
        ("Worst ball (total void)", f"{summary.get('maximum_void_ratio_percent', 0.0)} %"),
        ("Worst single void", f"{summary.get('max_single_void_ratio_percent', 0.0)} %"),
        (
            "Assessment split",
            f"{counts['PASS']} acceptable / {counts['REVIEW']} process "
            f"indicator / {counts['FAIL']} defect",
        ),
    ]
    table = document.add_table(rows=len(metric_pairs), cols=2)
    table.style = "Light Grid Accent 1"
    for idx, (label, value) in enumerate(metric_pairs):
        cells = table.rows[idx].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value

    # Histogram.
    histogram = _histogram_png(rows)
    if histogram is not None:
        document.add_picture(histogram, width=Inches(6.0))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # IPC criteria table.
    document.add_heading("IPC criteria evaluation", level=1)
    criteria_rows = _criteria_rows(summary, board, criteria, context)
    table = document.add_table(rows=len(criteria_rows) + 1, cols=5)
    table.style = "Light Grid Accent 1"
    for col, header in enumerate(
        ["Criterion", "Source", "Limit", "Measured", "Result"]
    ):
        cell = table.rows[0].cells[col]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        _shade_cell(cell, COLOR_HEADER_BG)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, (crit, source, limit, measured, result) in enumerate(
        criteria_rows, start=1
    ):
        values = [crit, source, limit, measured, result]
        for c_idx, value in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run_item in paragraph.runs:
                    run_item.font.size = Pt(9)
        result_cell = table.rows[r_idx].cells[4]
        if result.startswith("FAIL"):
            _shade_cell(result_cell, COLOR_FAIL_LIGHT)
        elif result in ("INFO", "REVIEW"):
            _shade_cell(result_cell, COLOR_REVIEW_LIGHT)
        elif result == "PASS":
            _shade_cell(result_cell, COLOR_PASS_LIGHT)

    # Worst balls.
    document.add_heading("Worst 10 balls by total void", level=1)
    top = sorted(
        rows,
        key=lambda r: float(r.get("void_ratio_percent", 0.0)),
        reverse=True,
    )[:10]
    table = document.add_table(rows=len(top) + 1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = [
        "Ball ID", "X (px)", "Y (px)", "Total void (%)",
        "Largest void (%)", "IPC assessment",
    ]
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    for r_idx, row in enumerate(top, start=1):
        values = [
            str(row["ball_id"]),
            str(row["center_x"]),
            str(row["center_y"]),
            f"{float(row['void_ratio_percent']):.2f}",
            f"{float(row['largest_void_ratio_percent']):.2f}",
            str(row["ipc_assessment"]),
        ]
        for c_idx, value in enumerate(values):
            table.rows[r_idx].cells[c_idx].text = value
        quality = str(row.get("quality", ""))
        if quality == "FAIL":
            _shade_cell(table.rows[r_idx].cells[5], COLOR_FAIL_LIGHT)
        elif quality == "REVIEW":
            _shade_cell(table.rows[r_idx].cells[5], COLOR_REVIEW_LIGHT)

    # Annotated preview.
    preview = context.get("preview_image")
    if preview and Path(str(preview)).exists():
        document.add_heading("Annotated X-ray (detected voids)", level=1)
        document.add_picture(str(preview), width=Inches(6.3))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Honest analysis.
    document.add_heading("Algorithm analysis - honest assessment", level=1)
    for section, bullets in _honest_analysis_sections(rows, summary, context):
        document.add_heading(section, level=2)
        for bullet in bullets:
            document.add_paragraph(bullet, style="List Bullet")

    footer = document.add_paragraph(
        f"Generated automatically by {TOOL_NAME} on {generated}. "
        "Verdicts follow IPC-A-610 (25% cumulative projected void area per "
        "ball) with IPC-7095 process guidance; they are based on a single "
        "top-down 2D X-ray projection and should be confirmed by an "
        "operator for boards flagged REVIEW or REJECT."
    )
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(8)

    document.save(str(path))
